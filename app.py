import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import plotly.io as pio
from sklearn.metrics import roc_curve, auc, confusion_matrix
from docx import Document
from io import BytesIO

# =====================================================
# PAGE CONFIG
# =====================================================
st.set_page_config(
    page_title="NovaMart Marketing Analytics Dashboard",
    page_icon="📊",
    layout="wide"
)

# =====================================================
# 🔒 LOCKED HYBRID THEME (DO NOT TOUCH)
# =====================================================
st.markdown("""
<style>
.stApp {
    background-color: #0E1117;
    color: #FFFFFF;
}
h1,h2,h3,h4,h5,h6,p,span,label {
    color: #FFFFFF !important;
}
[data-testid="stSidebar"] {
    background-color: #1F4E79;
}
</style>
""", unsafe_allow_html=True)

pio.templates.default = "plotly_dark"

# =====================================================
# DATA LOADING
# =====================================================
@st.cache_data
def load_csv(file):
    return pd.read_csv(file)

campaign_df = load_csv("campaign_performance.csv")
customer_df = load_csv("customer_data.csv")
learning_curve_df = load_csv("learning_curve.csv")
funnel_df = load_csv("funnel_data.csv")
geo_df = load_csv("geographic_data.csv")

customer_df.columns = customer_df.columns.str.lower()
learning_curve_df.columns = learning_curve_df.columns.str.lower()

# =====================================================
# GLOBAL LTV COMPUTATION (CRITICAL)
# =====================================================
customer_df["safe_churn"] = customer_df["churn_probability"].replace(0, 0.001)
customer_df["ltv"] = customer_df["avg_order_value"] / customer_df["safe_churn"]

# =====================================================
# SIDEBAR NAVIGATION
# =====================================================
st.sidebar.title("📌 Navigation")
page = st.sidebar.radio(
    "Go to",
    [
        "Executive Overview",
        "Campaign Analytics",
        "Customer Insights",
        "Geographic Analysis",
        "ML Model Evaluation",
        "📄 Export Insights"
    ]
)

# =====================================================
# PAGE 1 — EXECUTIVE OVERVIEW
# =====================================================
if page == "Executive Overview":

    st.title("📊 NovaMart Marketing Dashboard")

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Total Revenue", f"₹{campaign_df['revenue'].sum():,.0f}")
    c2.metric("Conversions", f"{campaign_df['conversions'].sum():,.0f}")
    c3.metric("Avg ROAS", f"{campaign_df['roas'].mean():.2f}")
    c4.metric("Customers", f"{len(customer_df):,}")

    revenue_ts = campaign_df.groupby("date", as_index=False)["revenue"].sum()
    revenue_ts["date"] = pd.to_datetime(revenue_ts["date"])

    fig = px.line(revenue_ts, x="date", y="revenue", title="Revenue Trend")
    st.plotly_chart(fig, use_container_width=True)

# =====================================================
# PAGE 2 — CAMPAIGN ANALYTICS
# =====================================================
elif page == "Campaign Analytics":

    st.title("📣 Campaign Analytics")

    channel_rev = campaign_df.groupby("channel", as_index=False)["revenue"].sum()
    fig = px.bar(channel_rev, x="revenue", y="channel", orientation="h",
                 title="Revenue by Channel")
    st.plotly_chart(fig, use_container_width=True)

# =====================================================
# PAGE 3 — CUSTOMER INSIGHTS
# =====================================================
elif page == "Customer Insights":

    st.title("👥 Customer Insights")

    # LTV by Segment
    ltv_seg = customer_df.groupby("customer_segment", as_index=False)["ltv"].mean()
    fig1 = px.bar(ltv_seg, x="customer_segment", y="ltv",
                  title="Lifetime Value (LTV) by Segment")
    st.plotly_chart(fig1, use_container_width=True)

    # Income vs LTV
    fig2 = px.scatter(
        customer_df,
        x="income",
        y="ltv",
        color="customer_segment",
        trendline="ols",
        title="Income vs Lifetime Value (LTV)"
    )
    st.plotly_chart(fig2, use_container_width=True)

# =====================================================
# PAGE 4 — GEOGRAPHIC ANALYSIS (FIXED)
# =====================================================
elif page == "Geographic Analysis":

    st.title("🗺 Geographic Performance (India)")

    fig = px.scatter_geo(
        geo_df,
        lat="latitude",
        lon="longitude",
        size="revenue",
        color="satisfaction",
        hover_name="state",
        projection="natural earth",
        title="State-wise Revenue & Satisfaction"
    )
    fig.update_geos(fitbounds="locations", visible=False)
    st.plotly_chart(fig, use_container_width=True)

# =====================================================
# PAGE 5 — ML MODEL EVALUATION
# =====================================================
elif page == "ML Model Evaluation":

    st.title("🤖 ML Model Diagnostics")

    lc = learning_curve_df.rename(
        columns={"training_size": "train_size", "validation_score": "val_score"}
    )

    fig = go.Figure()
    fig.add_trace(go.Scatter(x=lc["train_size"], y=lc["train_score"],
                             mode="lines+markers", name="Train"))
    fig.add_trace(go.Scatter(x=lc["train_size"], y=lc["val_score"],
                             mode="lines+markers", name="Validation"))

    fig.update_layout(title="Learning Curve",
                      xaxis_title="Training Size",
                      yaxis_title="Score")
    st.plotly_chart(fig, use_container_width=True)

# =====================================================
# PAGE 6 — EXPORT INSIGHTS AS DOCX
# =====================================================
elif page == "📄 Export Insights":

    st.title("📄 Export Board-Level Insights")

    def generate_doc():
        doc = Document()
        doc.add_heading("NovaMart Marketing Analytics – Insights Report", 0)

        doc.add_heading("1. Executive Overview", level=1)
        doc.add_paragraph(
            "NovaMart shows strong revenue growth driven primarily by Google Ads and Email. "
            "Overall ROAS remains healthy, indicating efficient marketing spend."
        )

        doc.add_heading("2. Customer Insights", level=1)
        doc.add_paragraph(
            "Premium customers deliver the highest Lifetime Value. "
            "Income positively correlates with LTV, though several mid-income customers show "
            "upgrade potential."
        )

        doc.add_heading("3. Geographic Performance", level=1)
        doc.add_paragraph(
            "Metro regions dominate revenue contribution. Certain high-revenue states show "
            "lower satisfaction, highlighting operational improvement opportunities."
        )

        doc.add_heading("4. ML Model Evaluation", level=1)
        doc.add_paragraph(
            "The learning curve indicates stable generalization performance. "
            "Additional training data may provide marginal gains."
        )

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    if st.button("📥 Export Insights as DOCX"):
        file = generate_doc()
        st.download_button(
            label="Download Insights Report",
            data=file,
            file_name="NovaMart_Marketing_Insights.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# =====================================================
# FOOTER
# =====================================================
st.markdown("<hr><center>NovaMart | Marketing Analytics Dashboard</center>",
            unsafe_allow_html=True)
